"""Generate the printable supplement pillbox guide linked from README.md."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Supplement-Pillbox-Guide.docx"
PILLBOX_URL = "https://gymbeam.com/adjustable-pillbox-gymbeam.html"


COLORS = {
    "ink": "162236",
    "muted": "536176",
    "line": "C9D3DF",
    "empty": "EFF3F7",
    "empty_ink": "7B8797",
    "morning": "176B87",
    "morning_fill": "E8F5F8",
    "lunch": "B66B14",
    "lunch_fill": "FFF3DF",
    "before": "8A3FFC",
    "before_fill": "F2ECFF",
    "after": "087F5B",
    "after_fill": "E7F7F0",
    "evening": "344E8C",
    "evening_fill": "EBF0FB",
    "rx": "9D2449",
    "rx_fill": "FBE9F0",
    "warning": "A64B00",
    "warning_fill": "FFF0D8",
    "white": "FFFFFF",
}


COMPACT_QUANTITY_RE = re.compile(
    r"(?<=\d)\s+(?=(?:bln|mln|mcg|µg|μg|mg|kg|g|mL|ml|L|kcal|kJ|IU|CFU|"
    r"scoops?|capsules?|pills?|tablets?|softgels?|servings?)\b)",
    flags=re.IGNORECASE,
)


TIER_BY_CODE = {
    "M.1": "🥇",
    "M.2": "🥈",
    "M.4": "🥈",
    "M.5": "🥈",
    "M.6": "🥈",
    "M.8": "🥈",
    "M.9": "🥈",
    "M.10": "🥉",
    "M.11": "🥈",
    "M.13": "🥈",
    "M.14": "🥉",
    "M.15": "🥉",
    "M.18": "🥉",
    "L.2": "🥇",
    "L.3": "🥇",
    "L.4": "🥇",
    "L.5": "🥈",
    "L.6": "🥈",
    "L.7": "🥈",
    "L.9": "🥈",
    "L.10": "🥈",
    "L.11": "🥈",
    "L.13": "🥉",
    "L.14": "🥈",
    "L.15": "🥈",
    "L.16": "🥈",
    "L.17": "🥈",
    "L.18": "🥉",
    "E.3": "🥈",
    "E.4": "🥈",
    "E.5": "🥈",
    "E.6": "🥈",
    "E.7": "🥈",
    "E.8": "🥈",
    "E.9": "🥈",
    "E.10": "🥈",
    "E.11": "🥈",
    "E.12": "🥈",
    "E.13": "🥈",
    "E.15": "🥉",
    "E.16": "🥉",
    "E.17": "🥉",
    "E.18": "🥉",
    "B.1": "🥈",
    "B.2": "🥈",
    "B.3": "🥈",
    "B.4": "🥈",
    "B.5": "🥈",
    "B.6": "🥉",
    "B.7": "🥉",
    "B.8": "🥉",
    "B.9": "🥉",
    "A.1": "🥈",
    "A.2": "🥈",
    "A.3": "🥈",
    "A.4": "🥈",
    "A.9": "🥈",
}


def compact_quantity_spacing(value: str) -> str:
    """Keep numeric quantities visually compact, e.g. 250mg and 1pill."""
    return COMPACT_QUANTITY_RE.sub("", value)


def item(code, name, dose="", qty="", icons="", note="", kind="active", same_slot_count=1):
    return {
        "code": code,
        "name": compact_quantity_spacing(name),
        "dose": compact_quantity_spacing(dose),
        "qty": compact_quantity_spacing(qty),
        "icons": icons,
        "note": compact_quantity_spacing(note),
        "kind": kind,
        "same_slot_count": same_slot_count,
        "tier": TIER_BY_CODE.get(code, ""),
    }


def empty(code):
    return item(code, "EMPTY", qty="Available compartment", kind="empty")


PAGES = [
    {
        "title": "MORNING • BOX 1",
        "range": "M.1–M.9",
        "accent": "morning",
        "callout": "Breakfast pill organizer • Row-major placement: left to right, top to bottom.",
        "cells": [
            item("M.1", "Bifidobacterium longum 35624®", "1 bln CFU", icons="🦠 🚽"),
            item(
                "M.2",
                "Limosilactobacillus reuteri Gastrus®",
                "200 mln CFU",
                icons="🦠 🚽",
                note="Lactobacillus reuteri DSM 17938 + Lactobacillus reuteri ATCC PTA 6475",
            ),
            empty("M.3"),
            item(
                "M.4",
                "CoQ10 (Ubiquinone)",
                "100 mg",
                icons="❤️ 🧠 🛡️",
                note="Heat-dispersed soy-oil softgel • riboflavin 1.4 mg",
            ),
            item("M.5", "Acetyl-L-Carnitine", "750 mg", icons="🧠 ⚡ 🔄"),
            item("M.6", "Acetyl-L-Carnitine", "750 mg", icons="🧠 ⚡ 🔄"),
            empty("M.7"),
            item("M.8", "Silicon", "14 mg", icons="💇‍♂️ 👨‍🦳 🦴"),
            item(
                "M.9",
                "Lion's Mane Fruiting-Body Extract",
                "500 mg",
                icons="🧠 🛡️ 🔄",
                note="Beta-(1,3)(1,6)-Glucans 300 mg (30%)",
                same_slot_count=2,
            ),
        ],
    },
    {
        "title": "MORNING • BOX 2",
        "range": "M.10–M.18",
        "accent": "morning",
        "callout": "M.15 holds all three small NR capsules together.",
        "cells": [
            item("M.10", "L-Ergothioneine", "25 mg", icons="🧠 👨‍🦳 🛡️ 🔄", note="ErgoActive® • Vitamin C 5 mg"),
            item("M.11", "NAC", "1 g", icons="🛡️ 🫁 🔄"),
            empty("M.12"),
            item("M.13", "Astaxanthin", "18 mg", icons="👁️ 🛡️ 👨‍🦳"),
            item(
                "M.14",
                "Vitamin K1 + K2 MK-4 + K2 MK-7",
                "K1 1.5 mg • K2 MK-4 1 mg\nK2 MK-7 0.2 mg",
                icons="🦴 ❤️",
            ),
            item(
                "M.15",
                "Nicotinamide Riboside",
                "300 mg",
                icons="🧠 ⚡",
                same_slot_count=3,
            ),
            empty("M.16"),
            empty("M.17"),
            item("M.18", "Spermidine HCl", "50 mg (Spermidine 28 mg)", icons="⏰ 🧠"),
        ],
    },
    {
        "title": "LUNCH • BOX 1",
        "range": "L.1–L.9",
        "accent": "lunch",
        "callout": "Omega-3 is conditional: use L.2–L.4 on non-Fatty-Fish days only.",
        "cells": [
            {
                "code": "L.1",
                "name": "PRESCRIPTIONS",
                "kind": "shared",
                "parts": [
                    item("L.1.1", "Ezetimibe", "10 mg", icons="❤️ ⚕️", kind="rx"),
                    item("L.1.2", "Tadalafil", "5 mg", icons="🍆 ❤️ 🫀 🍅 ⚕️", kind="rx"),
                ],
            },
            item("L.2", "Omega 3", "500 mg EPA • 250 mg DHA", icons="🧠 ❤️ 👁️ 🔥", note="NON-FATTY-FISH DAYS ONLY", kind="warning"),
            item("L.3", "Omega 3", "500 mg EPA • 250 mg DHA", icons="🧠 ❤️ 👁️ 🔥", note="NON-FATTY-FISH DAYS ONLY", kind="warning"),
            item("L.4", "Omega 3", "500 mg EPA • 250 mg DHA", icons="🧠 ❤️ 👁️ 🔥", note="NON-FATTY-FISH DAYS ONLY", kind="warning"),
            item("L.5", "Sodium Butyrate", "500 mg", icons="🦠 🚽", note="Butyric acid 400 mg"),
            item("L.6", "Cocoa Flavanols", "250 mg", icons="❤️ 🧠"),
            item("L.7", "Cocoa Flavanols", "250 mg", icons="❤️ 🧠"),
            empty("L.8"),
            item("L.9", "Lycopene", "20 mg", icons="🍅 ❤️ 👨‍🦳 🛡️"),
        ],
    },
    {
        "title": "LUNCH • BOX 2",
        "range": "L.10–L.18",
        "accent": "lunch",
        "callout": "Evidence-weighted order continues from Lunch Box 1.",
        "cells": [
            item(
                "L.10",
                "Curcumin + Ginger + Turmerones",
                "Curcuminoids 200mg",
                icons="🔥 🧠 🦴 ❤️ 🚽 🍆 🔄",
                note="Advanced Curcumin Elite™ • Gingerols 60mg • Turmerones 75mg",
            ),
            item("L.11", "Ceratiq® Wheat Oil Extract", "350 mg", icons="👨‍🦳 💧 🔄", note="Phytoceramides • glycosylceramides • glycolipids"),
            empty("L.12"),
            item("L.13", "Broccoli Seed Extract", "200 mg", icons="🛡️ 🍅 🔥", note="Glucoraphanin 20 mg • Myrosinase"),
            item("L.14", "Berberine HCl", "490 mg", icons="🩸 🚽"),
            item("L.15", "Aged Black-Garlic Extract", "500 mg", icons="❤️ 🛡️ 🔄", note="SAC 2.5 mg • standardized to 0.5%"),
            item("L.16", "Phosphatidylserine", "300 mg", icons="🧠 😌"),
            item("L.17", "Milk Thistle", "380 mg", icons="🛡️ 🚽 🔄"),
            item("L.18", "DIM", "200 mg", icons="🛡️ 🍆 🔄"),
        ],
    },
    {
        "title": "EVENING • BOX 1",
        "range": "E.1–E.9",
        "accent": "evening",
        "callout": "Follow the frequency labels exactly.",
        "cells": [
            {
                "code": "E.1",
                "name": "PRESCRIPTIONS",
                "kind": "shared",
                "parts": [
                    item("E.1.1", "Isotretinoin", "10 mg", icons="👨‍🦳 ⚕️", note="WEEKLY ONLY", kind="rx"),
                    item("E.1.2", "Dutasteride", "0.5 mg", icons="💇‍♂️ 🍅 ⚕️", note="EVERY 2ND DAY", kind="rx"),
                ],
            },
            item("E.2", "Minoxidil", "5 mg", icons="💇‍♂️ 🧔 ⚕️", kind="rx"),
            item(
                "E.3",
                "Limosilactobacillus reuteri Gastrus®",
                "200 mln CFU",
                icons="🦠 🚽",
                note="Lactobacillus reuteri DSM 17938 + Lactobacillus reuteri ATCC PTA 6475",
            ),
            item("E.4", "Sodium Butyrate", "500 mg", icons="🦠 🚽", note="Butyric acid 400 mg"),
            item("E.5", "Melatonin + Botanical Extract Blend", "1 mg", icons="😴 😌", note="Passionflower • lemon balm • hops • saffron • Vitamin B6"),
            item("E.6", "L-Theanine", "400 mg", icons="😴 😌 🧠"),
            item("E.7", "Ashwagandha KSM-66", "500 mg", icons="😌 😴", note="Withanolides 25 mg"),
            item("E.8", "Glycine", "1,000 mg", icons="😴 🧠"),
            item("E.9", "Glycine", "1,000 mg", icons="😴 🧠"),
        ],
    },
    {
        "title": "EVENING • BOX 2",
        "range": "E.10–E.18",
        "accent": "evening",
        "callout": "Glycine continues from Evening Box 1.",
        "cells": [
            item("E.10", "Glycine", "1,000 mg", icons="😴 🧠"),
            item("E.11", "Berberine HCl", "490 mg", icons="🩸 🚽"),
            item("E.12", "Magnesium (Glycinate)", "133.3 mg", icons="💪 🧠 ❤️ 🦴"),
            item("E.13", "Milk Thistle", "380 mg", icons="🛡️ 🚽 🔄"),
            empty("E.14"),
            item("E.15", "GABA", "250 mg", icons="😴 😌 🔄", note="Magnesium citrate 20 mg"),
            item("E.16", "Inositol", "1,000 mg", icons="🧠 😌"),
            item("E.17", "Inositol", "1,000 mg", icons="🧠 😌"),
            item("E.18", "Apigenin", "200 mg", icons="😴 😌"),
        ],
    },
    {
        "title": "BEFORE WORKOUT",
        "range": "B.1–B.9",
        "accent": "before",
        "callout": "Take with the normal pre-workout drink. Powders use B.0 codes and are not pillbox compartments.",
        "cells": [
            item("B.1", "Pycnogenol", "100 mg", icons="🫀 🦴 👨‍🦳 ❤️", note="French maritime pine • OPC 65 mg"),
            item(
                "B.2",
                "Rhodiola rosea Root Extract",
                "500 mg",
                icons="🧠 😌 ⚡ 🍆",
                note="Rosavins 15 mg (3%) • Salidroside 5 mg (1%)",
            ),
            item("B.3", "Taurine", "1.5 g", icons="❤️ 💪 🫁"),
            item("B.4", "Raw Maca Root 6:1 Concentrate", "750 mg", icons="🍆 ⚡ 🔄", note="Gelatinized • 4.5g fresh-root equivalent"),
            item(
                "B.5",
                "Fenugreek Seed Extract",
                "300 mg",
                icons="🍆 🩸 🚽",
                note="Testofen® • Fenuside™ saponins 150mg (50%)",
                same_slot_count=2,
            ),
            item("B.6", "Tribulus Terrestris", "1,500 mg", icons="🍆"),
            item("B.7", "Calcium Alpha-Ketoglutarate (CaAKG)", "500 mg", icons="⏰ 💪"),
            item("B.8", "Calcium Alpha-Ketoglutarate (CaAKG)", "500 mg", icons="⏰ 💪"),
            item("B.9", "Calcium Alpha-Ketoglutarate (CaAKG)", "500 mg", icons="⏰ 💪"),
        ],
    },
    {
        "title": "AFTER WORKOUT",
        "range": "A.1–A.9",
        "accent": "after",
        "callout": "Swallow these pills first; then start the A.0.1 WPI + A.0.2 EAA drink. Empty cells are intentional.",
        "cells": [
            item("A.1", "UC-II® Type II Collagen", "40 mg", icons="🦴 💪 🔄", note="Standardized chicken cartilage • Total collagen 10 mg"),
            item("A.2", "Urolithin A", "500 mg", icons="⏰ 💪"),
            item("A.3", "MSM (OptiMSM®)", "1,500 mg", icons="🦴 🔥 👨‍🦳 🔄"),
            item("A.4", "MSM (OptiMSM®)", "1,500 mg", icons="🦴 🔥 👨‍🦳 🔄"),
            empty("A.5"),
            empty("A.6"),
            empty("A.7"),
            empty("A.8"),
            item("A.9", "Glucosamine Sulfate 2KCl", "1,400 mg", icons="🦴 🛡️"),
        ],
    },
]


def shade(element, fill: str):
    tc_pr = element._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=95, start=115, bottom=85, end=115):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color="C9D3DF", size="10"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        edge_el = borders.find(qn(tag))
        if edge_el is None:
            edge_el = OxmlElement(tag)
            borders.append(edge_el)
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), size)
        edge_el.set(qn("w:color"), color)


def set_repeatable_font(run, size, bold=False, color=None, name="Aptos"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def add_hyperlink(paragraph, text, url, color="176B87"):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    run.append(r_pr)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_repeatable_font(run, 7.5, color=COLORS["muted"])
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    run = paragraph.add_run(" / ")
    set_repeatable_font(run, 7.5, color=COLORS["muted"])
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " NUMPAGES "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def compact_paragraph(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def fill_simple_cell(cell, data, accent_key):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    kind = data.get("kind", "active")
    if kind == "empty":
        shade(cell, COLORS["empty"])
        border = COLORS["line"]
        accent = COLORS["empty_ink"]
    elif kind == "rx":
        shade(cell, COLORS["rx_fill"])
        border = COLORS["rx"]
        accent = COLORS["rx"]
    elif kind == "warning":
        shade(cell, COLORS["warning_fill"])
        border = COLORS["warning"]
        accent = COLORS["warning"]
    else:
        shade(cell, COLORS[f"{accent_key}_fill"])
        border = COLORS[accent_key]
        accent = COLORS[accent_key]
    set_cell_borders(cell, border, "11")

    p = cell.paragraphs[0]
    compact_paragraph(p, after=3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if data.get("tier"):
        run = p.add_run(f"{data['tier']} ")
        set_repeatable_font(run, 10.5, bold=True, color=accent, name="Segoe UI Emoji")
    run = p.add_run(data["code"])
    set_repeatable_font(run, 10.5, bold=True, color=accent)

    p = cell.add_paragraph()
    compact_paragraph(p, after=2, line=0.95)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    same_slot_count = data.get("same_slot_count", 1)
    display_name = f"{same_slot_count}x {data['name']}" if same_slot_count > 1 else data["name"]
    run = p.add_run(display_name)
    set_repeatable_font(run, 12.2 if kind != "empty" else 15, bold=True, color=COLORS["ink"] if kind != "empty" else COLORS["empty_ink"])

    if data.get("dose"):
        p = cell.add_paragraph()
        compact_paragraph(p, after=2, line=0.95)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data["dose"])
        set_repeatable_font(run, 9.3, bold=True, color=accent)

    if data.get("qty"):
        p = cell.add_paragraph()
        compact_paragraph(p, after=2, line=0.95)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data["qty"])
        set_repeatable_font(run, 8.2, color=COLORS["muted"])

    if data.get("icons"):
        p = cell.add_paragraph()
        compact_paragraph(p, after=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        icon_tokens = data["icons"].split()
        run = p.add_run(" ".join(icon_tokens))
        set_repeatable_font(run, 9 if len(icon_tokens) >= 5 else 10, name="Segoe UI Emoji")

    if data.get("note"):
        p = cell.add_paragraph()
        compact_paragraph(p, before=1, line=0.95)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data["note"])
        note_color = COLORS["warning"] if kind in ("warning", "rx") else COLORS["muted"]
        set_repeatable_font(run, 7.5, bold=kind in ("warning", "rx"), color=note_color)


def fill_shared_cell(cell, data, accent_key):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=75, start=70, bottom=65, end=70)
    shade(cell, COLORS["rx_fill"])
    set_cell_borders(cell, COLORS["rx"], "11")

    p = cell.paragraphs[0]
    compact_paragraph(p, after=1)
    run = p.add_run(data["code"])
    set_repeatable_font(run, 9.5, bold=True, color=COLORS["rx"])

    nested = cell.add_table(rows=1, cols=2)
    nested.alignment = WD_TABLE_ALIGNMENT.CENTER
    nested.autofit = False
    for idx, part in enumerate(data["parts"]):
        sub = nested.cell(0, idx)
        sub.width = Inches(1.47)
        fill_simple_cell(sub, part, accent_key)
        set_cell_margins(sub, top=60, start=45, bottom=55, end=45)


def add_page(document, page_data, page_index):
    accent_key = page_data["accent"]
    accent = COLORS[accent_key]

    title_table = document.add_table(rows=1, cols=2)
    title_table.autofit = False
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_table.columns[0].width = Inches(7.95)
    title_table.columns[1].width = Inches(2.1)
    left = title_table.cell(0, 0)
    right = title_table.cell(0, 1)
    for cell in (left, right):
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        set_cell_borders(cell, COLORS["white"], "0")
    left.text = ""
    p = left.paragraphs[0]
    compact_paragraph(p, after=0)
    run = p.add_run(page_data["title"])
    set_repeatable_font(run, 18, bold=True, color=COLORS["ink"])
    right.text = ""
    p = right.paragraphs[0]
    compact_paragraph(p, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(page_data["range"])
    set_repeatable_font(run, 16, bold=True, color=accent)

    p = document.add_paragraph()
    compact_paragraph(p, before=1, after=4)
    run = p.add_run(page_data["callout"])
    set_repeatable_font(run, 8.2, color=COLORS["muted"])

    table = document.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    width = Inches(3.31)
    for col in table.columns:
        col.width = width
    for row in table.rows:
        row.height = Inches(2.08)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for index, data in enumerate(page_data["cells"]):
        cell = table.cell(index // 3, index % 3)
        cell.width = width
        if data.get("kind") == "shared":
            fill_shared_cell(cell, data, accent_key)
        else:
            fill_simple_cell(cell, data, accent_key)

    p = document.add_paragraph()
    compact_paragraph(p, before=3, after=0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_hyperlink(p, "GymBeam Adjustable PillBox", PILLBOX_URL, color=accent)
    run = p.add_run(" • Full doses and evidence notes: README.md")
    set_repeatable_font(run, 7.3, color=COLORS["muted"])
    run = p.add_run(" • 🥇 strongest • 🥈 supportive • 🥉 exploratory")
    set_repeatable_font(run, 6.8, color=COLORS["muted"], name="Segoe UI Emoji")

    if page_index < len(PAGES) - 1:
        p = document.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.31)
    section.bottom_margin = Inches(0.30)
    section.left_margin = Inches(0.39)
    section.right_margin = Inches(0.39)
    section.header_distance = Inches(0.12)
    section.footer_distance = Inches(0.14)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    doc.core_properties.title = "Supplement Pillbox Guide"
    doc.core_properties.subject = "Physical slot map for the health protocol"
    doc.core_properties.author = "Health Protocol"
    doc.core_properties.keywords = "pillbox, supplements, organizer, protocol"

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    compact_paragraph(p)
    run = p.add_run("HEALTH PROTOCOL • PHYSICAL SLOT MAP • 26 AUG 2026")
    set_repeatable_font(run, 6.5, bold=True, color=COLORS["muted"])

    footer = section.footer
    p = footer.paragraphs[0]
    compact_paragraph(p)
    add_page_field(p)

    for index, page_data in enumerate(PAGES):
        add_page(doc, page_data, index)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
